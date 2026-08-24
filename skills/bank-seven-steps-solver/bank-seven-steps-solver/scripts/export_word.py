#!/usr/bin/env python3
"""
七步成诗问题解决助手 - Word 文档导出脚本

根据用户交互记录生成专业的 Word 报告文档。
- 字体：微软雅黑（Microsoft YaHei）
- 页面：窄边距（上下左右各 1.27cm）
- 设计：深蓝色标题 + 专业报告格式

使用方法:
    python3 export_word.py --data data.json --output "报告.docx"
"""

import os
import sys
import json
import argparse
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ═══════════════════════════════════════════════════
# 常量
# ═══════════════════════════════════════════════════
FONT_NAME = '微软雅黑'          # 中文字体
FONT_NAME_EN = 'Microsoft YaHei'  # 英文字体回退
NAVY = RGBColor(0x05, 0x1C, 0x2C)  # McKinsey 深蓝
ACCENT_BLUE = RGBColor(0x00, 0x6B, 0xA6)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MED_GRAY = RGBColor(0x66, 0x66, 0x66)
NARROW_MARGIN = Cm(1.27)  # 窄边距

# ═══════════════════════════════════════════════════
# 样式设置
# ═══════════════════════════════════════════════════

def setup_document():
    """创建并配置文档：窄边距 + 微软雅黑字体"""
    doc = Document()

    # 页面设置：窄边距
    for section in doc.sections:
        section.top_margin = NARROW_MARGIN
        section.bottom_margin = NARROW_MARGIN
        section.left_margin = NARROW_MARGIN
        section.right_margin = NARROW_MARGIN

    # 设置默认字体为微软雅黑
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME_EN
    font.size = Pt(11)
    font.color.rgb = DARK_GRAY
    # 设置东亚字体
    rPr = style.element.rPr
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}><w:rFonts w:eastAsia="{FONT_NAME}"/></w:rPr>')
        style.element.append(rPr)
    else:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
            rPr.append(rFonts)
        else:
            rFonts.set(qn('w:eastAsia'), FONT_NAME)

    # 配置标题样式
    _setup_heading_styles(doc)

    return doc


def _setup_heading_styles(doc):
    """配置标题样式"""
    # Heading 1 - 一级标题
    h1 = doc.styles['Heading 1']
    h1.font.name = FONT_NAME_EN
    h1.font.size = Pt(22)
    h1.font.color.rgb = NAVY
    h1.font.bold = True
    h1_rPr = h1.element.rPr
    if h1_rPr is not None:
        rFonts = h1_rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            rFonts.set(qn('w:eastAsia'), FONT_NAME)
        else:
            h1_rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>'))
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)

    # Heading 2 - 二级标题
    h2 = doc.styles['Heading 2']
    h2.font.name = FONT_NAME_EN
    h2.font.size = Pt(16)
    h2.font.color.rgb = ACCENT_BLUE
    h2.font.bold = True
    h2_rPr = h2.element.rPr
    if h2_rPr is not None:
        rFonts = h2_rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            rFonts.set(qn('w:eastAsia'), FONT_NAME)
        else:
            h2_rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>'))
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(8)

    # Heading 3 - 三级标题
    h3 = doc.styles['Heading 3']
    h3.font.name = FONT_NAME_EN
    h3.font.size = Pt(13)
    h3.font.color.rgb = DARK_GRAY
    h3.font.bold = True
    h3_rPr = h3.element.rPr
    if h3_rPr is not None:
        rFonts = h3_rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            rFonts.set(qn('w:eastAsia'), FONT_NAME)
        else:
            h3_rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>'))
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(6)


def _set_run_font(run, font_name=FONT_NAME, font_name_en=FONT_NAME_EN,
                  size=None, color=None, bold=False):
    """统一设置 run 的字体属性"""
    run.font.name = font_name_en
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    run.font.bold = bold
    # 设置东亚字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)


def add_paragraph(doc, text, style=None, alignment=None, bold=False,
                  font_size=None, font_color=None, space_before=None, space_after=None):
    """添加段落并设置格式"""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    _set_run_font(run, size=font_size, color=font_color, bold=bold)
    if alignment:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    return p


def add_colored_table(doc, headers, rows, col_widths=None):
    """添加带颜色的专业表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表头行 - 深蓝色背景
    hdr_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        _set_run_font(run, size=Pt(10), color=RGBColor(0xFF, 0xFF, 0xFF), bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 深蓝色背景
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="051C2C" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # 数据行
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, value in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            _set_run_font(run, size=Pt(10), color=DARK_GRAY)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # 交替行背景
            if r_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>')
                cell._element.get_or_add_tcPr().append(shading)

    # 设置列宽
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = width

    return table


def add_separator(doc):
    """添加分隔线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('─' * 60)
    _set_run_font(run, size=Pt(8), color=RGBColor(0xCC, 0xCC, 0xCC))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ═══════════════════════════════════════════════════
# 封面页
# ═══════════════════════════════════════════════════

def create_cover_page(doc, data):
    """生成封面页"""
    # 空行填充
    for _ in range(6):
        doc.add_paragraph()

    # 标题
    title = data.get('title', '七步成诗问题解决报告')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    _set_run_font(run, size=Pt(28), color=NAVY, bold=True)

    # 副标题
    subtitle = data.get('subtitle', '基于麦肯锡七步成诗法的系统化问题分析')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    _set_run_font(run, size=Pt(16), color=ACCENT_BLUE)

    # 分隔线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 30)
    _set_run_font(run, size=Pt(12), color=NAVY)

    # 日期
    date_str = data.get('date', datetime.now().strftime('%Y年%m月'))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run(date_str)
    _set_run_font(run, size=Pt(14), color=MED_GRAY)

    # 方法论说明
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run('基于麦肯锡七步成诗法（Seven Steps Problem Solving）')
    _set_run_font(run, size=Pt(10), color=MED_GRAY)

    # 分页
    doc.add_page_break()


# ═══════════════════════════════════════════════════
# 目录页
# ═══════════════════════════════════════════════════

def create_toc_page(doc):
    """生成目录页"""
    doc.add_heading('目录', level=1)
    doc.add_paragraph()

    toc_items = [
        ('一', 'Step 1: 界定问题（Problem Statement）'),
        ('二', 'Step 2: 分解问题（Logic Tree）'),
        ('三', 'Step 3: 优先排序（Priority Matrix）'),
        ('四', 'Step 4&5: 工作计划与关键分析'),
        ('五', 'Step 6: 归纳建议（Pyramid Principle）'),
        ('六', 'Step 7: 交流沟通（Storyboard & Elevator Pitch）'),
        ('七', '附录：交互过程摘要'),
    ]

    for num, title in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{num}、{title}')
        _set_run_font(run, size=Pt(13), color=NAVY, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# Step 1: 问题陈述
# ═══════════════════════════════════════════════════

def create_step1_problem_statement(doc, data):
    """Step 1: 问题陈述"""
    ps = data.get('problem_statement', {})

    doc.add_heading('Step 1: 界定问题', level=1)
    add_paragraph(doc, '基于 SMART 原则，对核心问题进行清晰界定。',
                  font_size=Pt(11), font_color=MED_GRAY, space_after=Pt(12))

    # 核心问题框
    doc.add_heading('核心问题', level=2)
    p = doc.add_paragraph()
    run = p.add_run(ps.get('core_problem', '待填写'))
    _set_run_font(run, size=Pt(14), color=NAVY, bold=True)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)

    # 五要素表格
    doc.add_heading('问题陈述表', level=2)

    fields = [
        ('1. 现状/背景', ps.get('background', '待填写')),
        ('2. 决策者与相关方', ps.get('stakeholders', '待填写')),
        ('3. 成功标准', ps.get('success_criteria', '待填写')),
        ('4. 解决方案范围', ps.get('scope', '待填写')),
        ('5. 限制因素', ps.get('constraints', '待填写')),
    ]

    add_colored_table(doc,
                      headers=['要素', '内容'],
                      rows=fields,
                      col_widths=[Cm(4), Cm(13)])

    # SMART 验证表
    doc.add_paragraph()
    doc.add_heading('SMART 验证', level=2)

    smart_items = [
        ('Specific（具体）', ps.get('smart_s', '✅ 问题足够具体')),
        ('Measurable（可衡量）', ps.get('smart_m', '✅ 成功标准可量化')),
        ('Actionable（可行动）', ps.get('smart_a', '✅ 指向可执行行动')),
        ('Relevant（相关）', ps.get('smart_r', '✅ 与核心业务相关')),
        ('Time-framed（有时限）', ps.get('smart_t', '✅ 有明确时限')),
    ]

    add_colored_table(doc,
                      headers=['SMART 原则', '检查结果'],
                      rows=smart_items,
                      col_widths=[Cm(5), Cm(12)])

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# Step 2: 逻辑树
# ═══════════════════════════════════════════════════

def create_step2_logic_tree(doc, data):
    """Step 2: 逻辑树"""
    tree = data.get('logic_tree', {})
    root = tree.get('root', '核心问题')
    branches = tree.get('branches', [])

    doc.add_heading('Step 2: 分解问题（逻辑树）', level=1)
    add_paragraph(doc, '基于 MECE 原则，将核心问题分解为可独立分析的子议题。',
                  font_size=Pt(11), font_color=MED_GRAY, space_after=Pt(12))

    # 根节点
    doc.add_heading(f'核心问题：{root}', level=2)

    # 逻辑树文本
    for i, branch in enumerate(branches):
        branch_name = branch.get('name', f'子议题{i+1}')
        is_last = (i == len(branches) - 1)
        prefix = '└── ' if is_last else '├── '

        p = doc.add_paragraph()
        run = p.add_run(f'{prefix}{branch_name}')
        _set_run_font(run, size=Pt(12), color=NAVY, bold=True)
        p.paragraph_format.left_indent = Cm(1)

        children = branch.get('children', [])
        child_prefix_str = '    ' if is_last else '│   '
        for j, child in enumerate(children):
            is_last_child = (j == len(children) - 1)
            connector = '└── ' if is_last_child else '├── '
            p2 = doc.add_paragraph()
            run2 = p2.add_run(f'{child_prefix_str}{connector}{child}')
            _set_run_font(run2, size=Pt(11), color=DARK_GRAY)
            p2.paragraph_format.left_indent = Cm(1)
            p2.paragraph_format.space_before = Pt(2)
            p2.paragraph_format.space_after = Pt(2)

    # MECE 检查
    doc.add_paragraph()
    doc.add_heading('MECE 验证', level=2)

    p = doc.add_paragraph()
    run = p.add_run('相互独立性：')
    _set_run_font(run, size=Pt(11), color=NAVY, bold=True)
    run2 = p.add_run(tree.get('mece_exclusive', '所有子议题之间无概念重叠'))
    _set_run_font(run2, size=Pt(11), color=DARK_GRAY)

    p = doc.add_paragraph()
    run = p.add_run('完全穷尽性：')
    _set_run_font(run, size=Pt(11), color=NAVY, bold=True)
    run2 = p.add_run(tree.get('mece_exhaustive', '所有子议题合计覆盖了问题的全部方面'))
    _set_run_font(run2, size=Pt(11), color=DARK_GRAY)

    # 分解维度表
    doc.add_paragraph()
    doc.add_heading('分解维度说明', level=3)

    tree_rows = []
    for i, branch in enumerate(branches):
        tree_rows.append((f'L1-{i+1}', branch.get('name', ''),
                          branch.get('logic', '按维度分解'), '✅'))
        for child in branch.get('children', []):
            tree_rows.append(('  L2', child, '—', '✅'))

    add_colored_table(doc,
                      headers=['层级', '议题', '分解逻辑', 'MECE'],
                      rows=tree_rows,
                      col_widths=[Cm(2), Cm(5), Cm(6), Cm(2)])

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# Step 3: 优先排序
# ═══════════════════════════════════════════════════

def create_step3_priority_matrix(doc, data):
    """Step 3: 优先排序矩阵"""
    matrix = data.get('priority_matrix', {})

    doc.add_heading('Step 3: 优先排序（2×2 矩阵）', level=1)
    add_paragraph(doc, '根据「影响重要性」和「可行性」筛选最优先解决的议题。',
                  font_size=Pt(11), font_color=MED_GRAY, space_after=Pt(12))

    # 四象限
    quadrants = [
        ('⭐ 最高优先级（高影响 × 高可行）', matrix.get('priority', []), '立即启动'),
        ('🟡 快速解决（低影响 × 高可行）', matrix.get('quick_win', []), '本周完成'),
        ('🔵 重点规划（高影响 × 低可行）', matrix.get('plan', []), '制定计划'),
        ('⚪ 暂时搁置（低影响 × 低可行）', matrix.get('shelve', []), '纳入观察'),
    ]

    for label, items, strategy in quadrants:
        doc.add_heading(label, level=2)
        if items:
            for item in items:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(item)
                _set_run_font(run, size=Pt(11), color=DARK_GRAY)
        else:
            add_paragraph(doc, '（无）', font_size=Pt(11), font_color=MED_GRAY)

        p = doc.add_paragraph()
        run = p.add_run(f'策略：{strategy}')
        _set_run_font(run, size=Pt(10), color=MED_GRAY)

    # 优先级汇总表
    doc.add_paragraph()
    doc.add_heading('优先级汇总', level=2)

    all_items = []
    priority_labels = [
        ('priority', '🔴 最高优先级'),
        ('quick_win', '🟡 快速解决'),
        ('plan', '🔵 重点规划'),
        ('shelve', '⚪ 暂时搁置'),
    ]
    for key, label in priority_labels:
        for item in matrix.get(key, []):
            all_items.append((item, label))

    if all_items:
        add_colored_table(doc,
                          headers=['议题', '优先级'],
                          rows=all_items,
                          col_widths=[Cm(10), Cm(5)])

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# Step 4&5: 工作计划与关键分析
# ═══════════════════════════════════════════════════

def create_step4_5_work_plan(doc, data):
    """Step 4&5: 工作计划与关键分析"""
    plan = data.get('work_plan', [])

    doc.add_heading('Step 4&5: 工作计划与关键分析', level=1)
    add_paragraph(doc, '为每个优先议题制定详细工作计划，设计关键分析方法。',
                  font_size=Pt(11), font_color=MED_GRAY, space_after=Pt(12))

    # 综合工作计划表
    doc.add_heading('综合工作计划表', level=2)

    if plan:
        rows = []
        for i, item in enumerate(plan):
            rows.append((
                str(i + 1),
                item.get('issue', ''),
                item.get('owner', ''),
                item.get('method', ''),
                item.get('deadline', ''),
                item.get('deliverable', ''),
                item.get('status', '待启动'),
            ))

        add_colored_table(doc,
                          headers=['序号', '议题', '负责人', '分析方法', '截止时间', '预期成果', '状态'],
                          rows=rows,
                          col_widths=[Cm(1), Cm(3.5), Cm(2), Cm(3), Cm(2), Cm(3), Cm(2)])
    else:
        add_paragraph(doc, '（待填写工作计划）', font_size=Pt(11), font_color=MED_GRAY)

    # 议题分析工作表
    doc.add_paragraph()
    doc.add_heading('议题分析工作表', level=2)

    for i, item in enumerate(plan):
        doc.add_heading(f'议题 {i+1}: {item.get("issue", "")}', level=3)

        fields = [
            ('初始假设', item.get('hypothesis', '待填写')),
            ('支持依据', item.get('supporting_evidence', '待填写')),
            ('分析方法', item.get('method', '待填写')),
            ('所需信息', item.get('required_info', '待填写')),
            ('信息来源', item.get('data_source', '待填写')),
            ('负责人', item.get('owner', '待填写')),
            ('时间节点', item.get('deadline', '待填写')),
            ('预期成果', item.get('deliverable', '待填写')),
        ]

        add_colored_table(doc,
                          headers=['维度', '内容'],
                          rows=fields,
                          col_widths=[Cm(3.5), Cm(13)])
        doc.add_paragraph()

    # 工作计划自检
    doc.add_heading('工作计划自检（麦肯锡 5 问）', level=2)

    checklist = [
        '目标和最终成果是否明确界定？',
        '所有分析是否都十分必要，能充分回答问题？',
        '下一步工作是否明确？分析是否切实可行？',
        '责任和时间要求是否明确？',
        '时间安排是否符合整体项目要求和重点？',
    ]

    for item in checklist:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'☐ {item}')
        _set_run_font(run, size=Pt(11), color=DARK_GRAY)

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# Step 6: 归纳建议（金字塔原理）
# ═══════════════════════════════════════════════════

def create_step6_pyramid(doc, data):
    """Step 6: 金字塔原理"""
    pyramid = data.get('pyramid', {})

    doc.add_heading('Step 6: 归纳建议（金字塔原理）', level=1)
    add_paragraph(doc, '使用金字塔原理自下而上归纳结论，构建严密的论证结构。',
                  font_size=Pt(11), font_color=MED_GRAY, space_after=Pt(12))

    # 中心思想
    doc.add_heading('中心思想', level=2)
    p = doc.add_paragraph()
    run = p.add_run(pyramid.get('central_idea', '待填写'))
    _set_run_font(run, size=Pt(14), color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(12)

    # 主线与论据
    mainlines = pyramid.get('mainlines', [])
    for i, ml in enumerate(mainlines):
        doc.add_heading(f'主线 {i+1}: {ml.get("title", "")}', level=2)

        evidence = ml.get('evidence', [])
        for ev in evidence:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(ev)
            _set_run_font(run, size=Pt(11), color=DARK_GRAY)

    # 金字塔检查清单
    doc.add_paragraph()
    doc.add_heading('金字塔质量检查', level=2)

    checks = {
        '中心思想': [
            '回答了决策者的关键问题',
            '是提炼而非信息罗列',
            '语言简洁准确',
        ],
        '主线': [
            '紧密、无重叠地支持中心思想（MECE）',
            '面向行动/解决方案',
            '属于同一层次且逻辑排列',
            '使用决策者熟悉的语言',
        ],
        '支持论据': [
            '相关、充分且基于事实',
            '充分支持上一层观点',
            '同一类别下属于同一逻辑类型',
        ],
    }

    for category, items in checks.items():
        doc.add_heading(category, level=3)
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(f'☐ {item}')
            _set_run_font(run, size=Pt(11), color=DARK_GRAY)

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# Step 7: 交流沟通
# ═══════════════════════════════════════════════════

def create_step7_communication(doc, data):
    """Step 7: 交流沟通"""
    pitch = data.get('elevator_pitch', {})
    storyboard = data.get('storyboard', [])

    doc.add_heading('Step 7: 交流沟通', level=1)
    add_paragraph(doc, '构建沟通故事板和 30 秒电梯推销，为正式汇报做准备。',
                  font_size=Pt(11), font_color=MED_GRAY, space_after=Pt(12))

    # 30 秒电梯推销
    doc.add_heading('30 秒电梯推销', level=2)

    p = doc.add_paragraph()
    run = p.add_run(f'沟通对象：{pitch.get("audience", "决策者")}')
    _set_run_font(run, size=Pt(11), color=MED_GRAY)

    add_separator(doc)

    # 核心结论
    p = doc.add_paragraph()
    run = p.add_run('核心结论：')
    _set_run_font(run, size=Pt(12), color=NAVY, bold=True)
    run2 = p.add_run(pitch.get('core_conclusion', ''))
    _set_run_font(run2, size=Pt(12), color=DARK_GRAY)

    # 关键发现
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('三个关键发现：')
    _set_run_font(run, size=Pt(12), color=NAVY, bold=True)

    for i, finding in enumerate(pitch.get('key_findings', [])):
        p = doc.add_paragraph()
        run = p.add_run(f'{i+1}. {finding}')
        _set_run_font(run, size=Pt(11), color=DARK_GRAY)
        p.paragraph_format.left_indent = Cm(1)

    # 建议行动
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('建议行动：')
    _set_run_font(run, size=Pt(12), color=NAVY, bold=True)
    run2 = p.add_run(pitch.get('next_action', ''))
    _set_run_font(run2, size=Pt(12), color=DARK_GRAY)

    add_separator(doc)

    p = doc.add_paragraph()
    run = p.add_run('⏱️ 预计时长：25-30 秒')
    _set_run_font(run, size=Pt(10), color=MED_GRAY)

    # 沟通故事板
    doc.add_paragraph()
    doc.add_heading('沟通故事板', level=2)

    add_paragraph(doc,
                  '故事线逻辑：情境（为什么）→ 冲突（问题是什么）→ 疑问（怎么办）→ 回答（我们的建议）',
                  font_size=Pt(11), font_color=ACCENT_BLUE, space_after=Pt(8))

    if storyboard:
        rows = []
        for i, item in enumerate(storyboard):
            rows.append((
                str(i + 1),
                item.get('title', ''),
                item.get('content_type', ''),
                item.get('chart_type', ''),
            ))

        add_colored_table(doc,
                          headers=['页码', '标题（故事线）', '内容类型', '核心数据/图表'],
                          rows=rows,
                          col_widths=[Cm(1.5), Cm(5), Cm(4), Cm(5)])

    doc.add_page_break()


# ═══════════════════════════════════════════════════
# 附录：交互过程摘要
# ═══════════════════════════════════════════════════

def create_appendix(doc, data):
    """附录：交互过程摘要"""
    doc.add_heading('附录：交互过程摘要', level=1)

    interaction_log = data.get('interaction_log', [])

    if interaction_log:
        for entry in interaction_log:
            step = entry.get('step', '')
            summary = entry.get('summary', '')

            if step:
                doc.add_heading(step, level=2)
            if summary:
                add_paragraph(doc, summary, font_size=Pt(11), font_color=DARK_GRAY)
    else:
        add_paragraph(doc,
                      '本报告由 AI 咨询顾问通过七步成诗法引导生成。'
                      '完整的交互过程记录请参阅同步导出的 Markdown 文件。',
                      font_size=Pt(11), font_color=MED_GRAY)

    # 报告生成信息
    doc.add_paragraph()
    add_separator(doc)
    add_paragraph(doc,
                  f'报告生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}',
                  font_size=Pt(9), font_color=MED_GRAY,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
                  '方法论：麦肯锡七步成诗法（Seven Steps Problem Solving）',
                  font_size=Pt(9), font_color=MED_GRAY,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)


# ═══════════════════════════════════════════════════
# 主函数
# ═══════════════════════════════════════════════════

def generate_word(data, output_path):
    """
    生成七步成诗问题解决报告 Word 文档。

    Args:
        data: dict 包含所有七步法的交互记录数据（与 export_ppt.py 相同的数据结构）
        output_path: str 输出文件路径

    Returns:
        str: 输出文件路径
    """
    doc = setup_document()

    # 生成各部分
    create_cover_page(doc, data)
    create_toc_page(doc)
    create_step1_problem_statement(doc, data)
    create_step2_logic_tree(doc, data)
    create_step3_priority_matrix(doc, data)
    create_step4_5_work_plan(doc, data)
    create_step6_pyramid(doc, data)
    create_step7_communication(doc, data)
    create_appendix(doc, data)

    # 确保输出目录存在
    outdir = os.path.dirname(output_path)
    if outdir:
        os.makedirs(outdir, exist_ok=True)

    # 保存
    doc.save(output_path)
    print(f"✅ Word 文档已生成: {output_path}")
    return output_path


def main():
    parser = argparse.ArgumentParser(description='七步成诗问题解决报告 Word 生成器')
    parser.add_argument('--output', '-o', default='七步成诗报告.docx', help='输出文件路径')
    parser.add_argument('--data', '-d', help='JSON 数据文件路径')
    args = parser.parse_args()

    if args.data and os.path.exists(args.data):
        with open(args.data, 'r', encoding='utf-8') as f:
            data = json.load(f)
    else:
        # 示例数据
        data = {
            "title": "七步成诗问题解决报告",
            "subtitle": "基于麦肯锡七步成诗法的系统化问题分析",
            "date": "2026年3月",
            "problem_statement": {
                "core_problem": "如何在2026年底前将信用卡业务市场份额提升3个百分点",
                "background": "信用卡市场竞争日趋激烈，互联网金融平台持续蚕食传统银行份额",
                "stakeholders": "决策者：零售银行部总经理；相关方：信用卡中心、科技部、风控部",
                "success_criteria": "市场份额从15%提升至18%，新增活卡量50万张",
                "scope": "仅限中国大陆市场，不含港澳台",
                "constraints": "年度营销预算不超过2亿元，需符合监管合规要求",
            },
            "logic_tree": {
                "root": "提升信用卡市场份额",
                "branches": [
                    {"name": "提升获客能力", "children": ["线上渠道获客", "线下网点获客", "合作伙伴渠道"]},
                    {"name": "提升客户价值", "children": ["提升交易频次", "提升客单价", "提升分期比例"]},
                    {"name": "降低客户流失", "children": ["提升满意度", "优化权益体系", "改善服务体验"]},
                ]
            },
            "priority_matrix": {
                "priority": ["线上渠道获客", "提升交易频次"],
                "quick_win": ["优化权益体系"],
                "plan": ["合作伙伴渠道"],
                "shelve": ["组织架构调整"],
            },
            "work_plan": [
                {"issue": "线上渠道获客", "owner": "张经理", "method": "数据分析+A/B测试",
                 "deadline": "Q2", "deliverable": "获客方案"},
            ],
            "pyramid": {
                "central_idea": "通过线上获客+交易频次提升双轮驱动",
                "mainlines": [
                    {"title": "线上获客能力是增长核心引擎",
                     "evidence": ["线上获客成本仅为线下1/3", "年轻客群90%来自线上"]},
                    {"title": "交易频次提升直接带动收入增长",
                     "evidence": ["频次提升10%对应收入增长8%"]},
                ]
            },
            "elevator_pitch": {
                "audience": "零售银行部总经理",
                "core_conclusion": "通过线上获客和交易频次双轮驱动实现份额提升",
                "key_findings": ["线上获客成本仅为线下1/3", "频次提升10%带动收入增长8%"],
                "next_action": "建议本周启动线上获客方案设计",
            },
            "storyboard": [
                {"title": "信用卡市场竞争态势", "content_type": "背景介绍", "chart_type": "趋势图"},
                {"title": "综合结论与建议", "content_type": "金字塔总结", "chart_type": "结构图"},
            ],
        }

    generate_word(data, args.output)


if __name__ == "__main__":
    main()
