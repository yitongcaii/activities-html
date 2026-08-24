#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文化活动SOP文档生成器
支持输出: Word (.docx), Excel (.xlsx), HTML (.html)
"""

import json
import sys
import os
import argparse
from datetime import datetime

# ============ Excel 输出 ============
def export_excel(content_dict, output_path):
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    except ImportError:
        print("缺少 openpyxl，正在安装...")
        os.system(f"{sys.executable} -m pip install openpyxl -q")
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    # Sheet 1: 活动定位与目标
    ws1 = wb.create_sheet("活动定位与目标")
    ws1.append(["活动类型", content_dict.get("activity_type", "")])
    ws1.append(["适用场景", content_dict.get("scenarios", "")])
    ws1.append([])
    ws1.append(["期望达成目标"])
    for level, goals in content_dict.get("goals", {}).items():
        ws1.append([level, goals])
    for row in ws1.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    # Sheet 2: 筹备时间线
    ws2 = wb.create_sheet("筹备时间线")
    ws2.append(["阶段", "事项", "负责人", "截止时间", "关键备注"])
    header_fill = PatternFill(start_color="F4A460", end_color="F4A460", fill_type="solid")
    for cell in ws2[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = header_fill
    for item in content_dict.get("timeline", []):
        ws2.append(item)
    for col in ws2.columns:
        max_len = 0
        col_letter = col[0].column_letter
        for cell in col:
            if cell.value:
                max_len = max(max_len, len(str(cell.value)))
        ws2.column_dimensions[col_letter].width = min(max_len + 4, 60)

    # Sheet 3: 通用清单
    ws3 = wb.create_sheet("通用筹备清单")
    ws3.append(["阶段", "事项", "负责人", "关键备注"])
    for cell in ws3[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = header_fill
    for item in content_dict.get("general_checklist", []):
        ws3.append(item)
    for col in ws3.columns:
        max_len = 0
        col_letter = col[0].column_letter
        for cell in col:
            if cell.value:
                max_len = max(max_len, len(str(cell.value)))
        ws3.column_dimensions[col_letter].width = min(max_len + 4, 60)

    # Sheet 4: 专项补充
    ws4 = wb.create_sheet("专项补充清单")
    ws4.append(["阶段", "补充事项", "负责人", "关键备注"])
    for cell in ws4[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = header_fill
    for item in content_dict.get("specific_checklist", []):
        ws4.append(item)
    for col in ws4.columns:
        max_len = 0
        col_letter = col[0].column_letter
        for cell in col:
            if cell.value:
                max_len = max(max_len, len(str(cell.value)))
        ws4.column_dimensions[col_letter].width = min(max_len + 4, 60)

    # Sheet 5: 风险提示
    ws5 = wb.create_sheet("风险提示")
    ws5.append(["风险项", "具体说明", "预防措施"])
    for cell in ws5[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = header_fill
    for item in content_dict.get("risks", []):
        ws5.append(item)
    for col in ws5.columns:
        max_len = 0
        col_letter = col[0].column_letter
        for cell in col:
            if cell.value:
                max_len = max(max_len, len(str(cell.value)))
        ws5.column_dimensions[col_letter].width = min(max_len + 4, 60)

    wb.save(output_path)
    print(f"Excel 已保存: {output_path}")


# ============ Word 输出 ============
def export_word(content_dict, output_path):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("缺少 python-docx，正在安装...")
        os.system(f"{sys.executable} -m pip install python-docx -q")
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 标题
    title = doc.add_heading(content_dict.get("activity_type", "文化活动筹备方案"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 活动定位
    doc.add_heading("一、活动定位与目标", level=1)
    doc.add_paragraph(f"活动类型: {content_dict.get('activity_type', '')}")
    doc.add_paragraph(f"适用场景: {content_dict.get('scenarios', '')}")
    doc.add_paragraph("期望达成目标:")
    for level, goals in content_dict.get("goals", {}).items():
        p = doc.add_paragraph(f"【{level}】{goals}", style="List Bullet")

    # 筹备时间线
    doc.add_heading("二、筹备时间线", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    headers = ["阶段", "事项", "负责人", "截止时间", "关键备注"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    for item in content_dict.get("timeline", []):
        row_cells = table.add_row().cells
        for i, val in enumerate(item):
            row_cells[i].text = val

    # 通用清单
    doc.add_heading("三、通用筹备清单", level=1)
    table2 = doc.add_table(rows=1, cols=4)
    table2.style = "Light Grid Accent 1"
    hdr_cells = table2.rows[0].cells
    for i, h in enumerate(["阶段", "事项", "负责人", "关键备注"]):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    for item in content_dict.get("general_checklist", []):
        row_cells = table2.add_row().cells
        for i, val in enumerate(item):
            row_cells[i].text = val

    # 专项补充
    doc.add_heading("四、专项补充清单", level=1)
    doc.add_paragraph(content_dict.get("specific_desc", ""))
    table3 = doc.add_table(rows=1, cols=4)
    table3.style = "Light Grid Accent 1"
    hdr_cells = table3.rows[0].cells
    for i, h in enumerate(["阶段", "补充事项", "负责人", "关键备注"]):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    for item in content_dict.get("specific_checklist", []):
        row_cells = table3.add_row().cells
        for i, val in enumerate(item):
            row_cells[i].text = val

    # 案例参考
    doc.add_heading("五、同类活动案例参考", level=1)
    for case in content_dict.get("cases", []):
        p = doc.add_paragraph()
        p.add_run(f"部门: {case['dept']}  |  阶段: {case['phase']}").bold = True
        doc.add_paragraph(f"活动: {case['activity']}")
        doc.add_paragraph()

    # 风险提示
    doc.add_heading("六、风险提示与注意事项", level=1)
    for risk in content_dict.get("risks", []):
        doc.add_paragraph(f"⚠️ {risk[0]} — {risk[1]}")
        if len(risk) > 2 and risk[2]:
            doc.add_paragraph(f"   预防: {risk[2]}", style="List Bullet 2")

    # 资源
    doc.add_heading("七、参考模板与资源", level=1)
    for r in content_dict.get("resources", []):
        doc.add_paragraph(r, style="List Bullet")

    doc.save(output_path)
    print(f"Word 已保存: {output_path}")


# ============ HTML 输出 ============
def export_html(content_dict, output_path):
    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{content_dict.get('activity_type', '文化活动筹备方案')}</title>
<style>
  body {{ font-family: 'Microsoft YaHei', sans-serif; margin: 40px; background: #fafafa; color: #333; }}
  .container {{ max-width: 960px; margin: 0 auto; background: #fff; padding: 40px; border-radius: 12px; box-shadow: 0 2px 12px rgba(0,0,0,0.08); }}
  h1 {{ color: #e67e22; border-bottom: 3px solid #e67e22; padding-bottom: 12px; text-align: center; }}
  h2 {{ color: #d35400; margin-top: 32px; border-left: 4px solid #e67e22; padding-left: 12px; }}
  h3 {{ color: #555; margin-top: 24px; }}
  table {{ width: 100%; border-collapse: collapse; margin: 16px 0; font-size: 14px; }}
  th {{ background: linear-gradient(135deg, #e67e22, #f39c12); color: #fff; padding: 12px; text-align: left; }}
  td {{ padding: 10px 12px; border-bottom: 1px solid #eee; }}
  tr:hover {{ background: #fdf5e6; }}
  .tag {{ display: inline-block; background: #e67e22; color: #fff; padding: 2px 10px; border-radius: 12px; font-size: 12px; margin-right: 6px; }}
  .risk {{ background: #fff5f5; border-left: 4px solid #e74c3c; padding: 12px; margin: 8px 0; border-radius: 4px; }}
  .case {{ background: #f0f8ff; border-left: 4px solid #3498db; padding: 12px; margin: 8px 0; border-radius: 4px; }}
  .resource {{ background: #f5fff5; border-left: 4px solid #27ae60; padding: 8px 12px; margin: 6px 0; border-radius: 4px; }}
  ul {{ line-height: 2; }}
  .highlight {{ color: #e67e22; font-weight: bold; }}
</style>
</head>
<body>
<div class="container">
<h1>{content_dict.get('activity_type', '文化活动筹备方案')}</h1>

<h2>一、活动定位与目标</h2>
<p><span class="tag">活动类型</span> {content_dict.get('activity_type', '')}</p>
<p><span class="tag">适用场景</span> {content_dict.get('scenarios', '')}</p>
<h3>期望达成目标</h3>
<ul>
"""
    for level, goals in content_dict.get("goals", {}).items():
        html += f"  <li><strong>【{level}】</strong>{goals}</li>\n"
    html += "</ul>\n"

    html += "<h2>二、筹备时间线</h2>\n<table>\n<tr><th>阶段</th><th>事项</th><th>负责人</th><th>截止时间</th><th>关键备注</th></tr>\n"
    for item in content_dict.get("timeline", []):
        html += f"<tr><td>{item[0]}</td><td>{item[1]}</td><td>{item[2]}</td><td>{item[3]}</td><td>{item[4]}</td></tr>\n"
    html += "</table>\n"

    html += "<h2>三、通用筹备清单</h2>\n<table>\n<tr><th>阶段</th><th>事项</th><th>负责人</th><th>关键备注</th></tr>\n"
    for item in content_dict.get("general_checklist", []):
        html += f"<tr><td>{item[0]}</td><td>{item[1]}</td><td>{item[2]}</td><td>{item[3]}</td></tr>\n"
    html += "</table>\n"

    html += f"<h2>四、{content_dict.get('activity_type', '')} 专项补充清单</h2>\n"
    html += f"<p>{content_dict.get('specific_desc', '')}</p>\n"
    html += "<table>\n<tr><th>阶段</th><th>补充事项</th><th>负责人</th><th>关键备注</th></tr>\n"
    for item in content_dict.get("specific_checklist", []):
        html += f"<tr><td>{item[0]}</td><td>{item[1]}</td><td>{item[2]}</td><td>{item[3]}</td></tr>\n"
    html += "</table>\n"

    html += "<h2>五、同类活动案例参考</h2>\n"
    for case in content_dict.get("cases", []):
        html += f"""<div class="case">
  <strong>部门:</strong> {case['dept']} &nbsp;|&nbsp; <strong>阶段:</strong> {case['phase']}<br>
  <strong>活动:</strong> {case['activity']}
</div>
"""

    html += "<h2>六、风险提示与注意事项</h2>\n"
    for risk in content_dict.get("risks", []):
        html += f"""<div class="risk">
  <strong>⚠️ {risk[0]}</strong><br>
  {risk[1]}<br>
  <em>预防: {risk[2] if len(risk) > 2 else '无'}</em>
</div>
"""

    html += "<h2>七、参考模板与资源</h2>\n"
    for r in content_dict.get("resources", []):
        html += f'<div class="resource">{r}</div>\n'

    html += """
</div>
</body>
</html>
"""
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html)
    print(f"HTML 已保存: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="文化活动SOP文档生成器")
    parser.add_argument("--format", choices=["word", "excel", "html", "all"], default="html",
                        help="输出格式")
    parser.add_argument("--input", required=True, help="输入JSON文件路径")
    parser.add_argument("--output", required=True, help="输出文件路径（不含扩展名，all模式会自动加扩展名）")
    args = parser.parse_args()

    with open(args.input, 'r', encoding='utf-8') as f:
        content = json.load(f)

    if args.format == "word" or args.format == "all":
        export_word(content, args.output + ".docx")
    if args.format == "excel" or args.format == "all":
        export_excel(content, args.output + ".xlsx")
    if args.format == "html" or args.format == "all":
        export_html(content, args.output + ".html")


if __name__ == "__main__":
    main()
