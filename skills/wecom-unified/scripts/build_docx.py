"""build_docx.py — Generate a .docx file from a JSONL spec.

Each line of the input file is a single command::

    {"action": "<function_name>", "params": {...}}

Workflow::

    1. 读 JSONL → 一次性按 ``references/doc-create.md`` 做参数校验
       （action 取值 / params 字段名 / 类型 / 取值范围）。
       任何偏差立即抛 ``TypeError``（"类型错误，无法执行"）。
    2. 校验通过后，再创建 docx 并按 action 派发到 ``DocxBuilder``。
    3. 通过本地文件系统写出 ``.docx``，输出路径自动选取于
       ``WECOMAGENT_WRITABLE_DIRS`` 的第一个目录；同名文件会追加
       ``_1`` / ``_2`` … 后缀避免覆盖。

Usage::

    python build_docx.py <spec.jsonl>
"""

from __future__ import annotations

import argparse
import base64
import functools
import io
import json
import os
import re
import sys
import time
from pathlib import Path
from typing import Any, Iterator, NamedTuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Emu, Pt, RGBColor

# ===========================================================================
# Constants & lookups
# ===========================================================================

_PARAGRAPH_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

_TABLE_ALIGN = {
    "left": WD_TABLE_ALIGNMENT.LEFT,
    "center": WD_TABLE_ALIGNMENT.CENTER,
    "right": WD_TABLE_ALIGNMENT.RIGHT,
}

EMU_PER_DXA = 635                  # 1 dxa = 1/20 pt = 635 EMU
DEFAULT_TABLE_TOTAL_DXA = 9072     # ~6.30 in, A4 content-area width
DEFAULT_LINE_DXA_NORMAL = 312
DEFAULT_LINE_DXA_HEADING = 408

# Table-level frame uses a thin theme-default line; per-cell borders
# use a soft gray, applied to every cell so the grid stays consistent
# on renderers that ignore table-level borders.
DEFAULT_TABLE_BORDER_COLOR_HEX = "auto"
DEFAULT_TABLE_BORDER_SIZE = 4
DEFAULT_CELL_BORDER_COLOR_HEX = "CBCDD1"
DEFAULT_CELL_BORDER_SIZE = 6

# --- XSD ordering anchors --------------------------------------------------
# Each tuple lists the children that the new element must appear *before*,
# per the OOXML schema. ``_set_unique_child`` inserts the new element ahead
# of the first sibling found.


def _anchors_after(tag: str, order: tuple) -> tuple:
    """Return the slice of ``order`` strictly after ``tag``."""
    return order[order.index(tag) + 1:]


# CT_PPrBase: shared by snapToGrid (pos 21) and spacing (pos 22) —
# every sibling listed comes after both.
_PPR_SPACING_ANCHORS = (
    qn("w:contextualSpacing"),
    qn("w:jc"),
    qn("w:outlineLvl"),
)

# CT_TblPr order (relevant prefix).
_TBL_PR_ORDER = (
    qn("w:tblW"),
    qn("w:jc"),
    qn("w:tblCellSpacing"),
    qn("w:tblInd"),
    qn("w:tblBorders"),
    qn("w:shd"),
    qn("w:tblLayout"),
    qn("w:tblLook"),
)
_TBL_PR_TBLW_ANCHORS = _anchors_after(qn("w:tblW"), _TBL_PR_ORDER)
_TBL_PR_BORDERS_ANCHORS = _anchors_after(qn("w:tblBorders"), _TBL_PR_ORDER)
_TBL_PR_LAYOUT_ANCHORS = _anchors_after(qn("w:tblLayout"), _TBL_PR_ORDER)

# CT_TcPrInner order (relevant prefix).
_TC_PR_ORDER = (
    qn("w:tcBorders"),
    qn("w:shd"),
    qn("w:noWrap"),
    qn("w:tcMar"),
    qn("w:textDirection"),
    qn("w:tcFitText"),
    qn("w:vAlign"),
    qn("w:hideMark"),
    qn("w:headers"),
)
_TC_PR_BORDERS_ANCHORS = _anchors_after(qn("w:tcBorders"), _TC_PR_ORDER)
_TC_PR_SHD_ANCHORS = _anchors_after(qn("w:shd"), _TC_PR_ORDER)
_TC_PR_MAR_ANCHORS = _anchors_after(qn("w:tcMar"), _TC_PR_ORDER)
_TC_PR_VALIGN_ANCHORS = _anchors_after(qn("w:vAlign"), _TC_PR_ORDER)


class _HeadingPreset(NamedTuple):
    style_name: str
    size_pt: int
    color_hex: str
    alignment: str | None


# Title 24pt → H1 18pt → H2 16pt → H3 14pt → H4 12pt → H5/H6 11pt.
_DEFAULT_HEADINGS: tuple[_HeadingPreset, ...] = (
    _HeadingPreset("Title",     24, "1A1A1A", "center"),
    _HeadingPreset("Subtitle",  18, "5C5C5C", "center"),
    _HeadingPreset("Heading 1", 18, "1A1A1A", None),
    _HeadingPreset("Heading 2", 16, "1A1A1A", None),
    _HeadingPreset("Heading 3", 14, "1A1A1A", None),
    _HeadingPreset("Heading 4", 12, "1A1A1A", None),
    _HeadingPreset("Heading 5", 11, "1A1A1A", None),
    _HeadingPreset("Heading 6", 11, "1A1A1A", None),
)

# 6 位十六进制颜色字符串，可选前缀 '#'。供 _hex_to_rgb / 上游校验复用。
_HEX_COLOR_RE = re.compile(r"^#?[0-9A-Fa-f]{6}$")


# ===========================================================================
# Exceptions
# ===========================================================================


class SpecTypeError(TypeError):
    """JSONL 参数校验失败抛出，等同 ``TypeError``，附带行号上下文。"""


# Backwards-compatible alias for any existing caller that imports SpecError.
SpecError = SpecTypeError


# ===========================================================================
# JSONL spec validation — 上游一次性校验（基于 references/doc-create.md）
# ===========================================================================
#
# 校验范围严格对齐 doc-create.md 中描述的 4 个 action 及其 params。
# 任何偏差均抛出 ``SpecTypeError``（继承自 ``TypeError``），由 main()
# 统一捕获并转成 "类型错误，无法执行" 提示。

# 段落 style 仅支持以下内置样式（doc-create.md：列表样式 + Subtitle）。
ALLOWED_PARAGRAPH_STYLES: frozenset[str] = frozenset({
    "List Bullet", "List Bullet 2", "List Bullet 3",
    "List Number", "List Number 2", "List Number 3",
    "Subtitle",
})

# 段落级对齐枚举。
ALLOWED_ALIGNMENTS: frozenset[str] = frozenset({
    "left", "center", "right", "justify",
})

# run / cell 对象支持的字段及类型（与 doc-create.md 中表格一致）。
# (int, float) 元组用于 "数字" 类（运行时显式排除 bool）。
RUN_FIELD_TYPES: dict[str, Any] = {
    "text": str,
    "bold": bool,
    "italic": bool,
    "underline": bool,
    "color_hex": str,
    "size_pt": (int, float),
    "font": str,
    "east_asia_font": str,
}

# add_heading.level 取值范围（doc-create.md：0=Title，1~4=章节标题）。
HEADING_LEVEL_MIN: int = 0
HEADING_LEVEL_MAX: int = 4

# 结构化输入上限：在校验阶段尽早拒绝异常输入，避免下游构建 / 序列化
MAX_COMMANDS: int = 5000              # 单份 JSONL 的命令条数上限
MAX_TEXT_CHARS: int = 20000           # 段落 text / run.text 单字段长度上限
MAX_RUNS_PER_PARAGRAPH: int = 200     # 单段 runs 数组长度上限
MAX_TABLE_ROWS: int = 10000            # 单表行数上限
MAX_TABLE_COLS: int = 50              # 单行列数上限
MAX_CELL_TEXT_CHARS: int = 5000       # 表格 cell（字符串或 run.text）长度上限


def _spec_raise(ctx: str, msg: str) -> None:
    raise SpecTypeError(f"{ctx}: {msg}")


def _spec_type_name(expected: Any) -> str:
    if isinstance(expected, type):
        return expected.__name__
    if isinstance(expected, tuple):
        return " | ".join(t.__name__ for t in expected if isinstance(t, type))
    return str(expected)


def _spec_check_type(value: Any, expected: Any, ctx: str, name: str) -> None:
    """对值做基础类型检查；显式拒绝 bool 充当 int/float。"""
    if expected is int:
        if isinstance(value, bool) or not isinstance(value, int):
            _spec_raise(ctx, f"参数 '{name}' 类型错误，期望 int，实际 {type(value).__name__}")
        return
    if expected is bool:
        if not isinstance(value, bool):
            _spec_raise(ctx, f"参数 '{name}' 类型错误，期望 bool，实际 {type(value).__name__}")
        return
    if isinstance(expected, tuple) and int in expected and float in expected:
        if isinstance(value, bool) or not isinstance(value, (int, float)):
            _spec_raise(ctx, f"参数 '{name}' 类型错误，期望 number，实际 {type(value).__name__}")
        return
    if not isinstance(value, expected):
        _spec_raise(
            ctx,
            f"参数 '{name}' 类型错误，期望 {_spec_type_name(expected)}，"
            f"实际 {type(value).__name__}",
        )


def _spec_check_hex_color(value: Any, ctx: str, name: str) -> None:
    if not (isinstance(value, str) and _HEX_COLOR_RE.match(value)):
        _spec_raise(
            ctx,
            f"参数 '{name}' 必须是 6 位十六进制颜色字符串"
            f"（如 'FF0000' 或 '#FF0000'），实际为 {value!r}",
        )


def _spec_validate_run_object(
    obj: Any,
    ctx: str,
    max_text_chars: int = MAX_TEXT_CHARS,
) -> None:
    """校验一个 run / table-cell 对象（字段集合相同）。

    ``max_text_chars`` 控制 ``text`` 字段的长度上限：段落里的 run 沿用
    ``MAX_TEXT_CHARS``；表格 cell 上下文则收窄到 ``MAX_CELL_TEXT_CHARS``。
    """
    if not isinstance(obj, dict):
        _spec_raise(ctx, f"必须是 JSON 对象（dict），实际 {type(obj).__name__}")

    unknown = set(obj) - set(RUN_FIELD_TYPES)
    if unknown:
        _spec_raise(
            ctx,
            f"包含未知字段 {sorted(unknown)}；允许字段: {sorted(RUN_FIELD_TYPES)}",
        )

    for fname, expected in RUN_FIELD_TYPES.items():
        if fname not in obj:
            continue
        _spec_check_type(obj[fname], expected, ctx, fname)

    if "text" in obj and len(obj["text"]) > max_text_chars:
        _spec_raise(
            ctx,
            f"参数 'text' 长度 {len(obj['text'])} 超过上限 {max_text_chars} 字符",
        )

    if "color_hex" in obj:
        _spec_check_hex_color(obj["color_hex"], ctx, "color_hex")

    if "size_pt" in obj:
        size = obj["size_pt"]
        if size < 1 or size > 819:
            _spec_raise(ctx, f"参数 'size_pt' 必须在 [1, 819] 范围内，实际为 {size}")


def _spec_validate_add_paragraph(params: dict, ctx: str) -> None:
    allowed = {"text", "runs", "style", "alignment"}
    unknown = set(params) - allowed
    if unknown:
        _spec_raise(
            ctx,
            f"add_paragraph 含未知参数 {sorted(unknown)}；允许参数: {sorted(allowed)}",
        )

    if "text" in params:
        _spec_check_type(params["text"], str, ctx, "text")
        if len(params["text"]) > MAX_TEXT_CHARS:
            _spec_raise(
                ctx,
                f"参数 'text' 长度 {len(params['text'])} 超过上限 "
                f"{MAX_TEXT_CHARS} 字符",
            )

    if "runs" in params:
        runs = params["runs"]
        if not isinstance(runs, list):
            _spec_raise(ctx, f"参数 'runs' 必须是数组，实际 {type(runs).__name__}")
        if len(runs) > MAX_RUNS_PER_PARAGRAPH:
            _spec_raise(
                ctx,
                f"参数 'runs' 数量 {len(runs)} 超过上限 "
                f"{MAX_RUNS_PER_PARAGRAPH}",
            )
        for i, r in enumerate(runs):
            _spec_validate_run_object(r, f"{ctx}.runs[{i}]")

    if "style" in params:
        style = params["style"]
        if not isinstance(style, str) or style not in ALLOWED_PARAGRAPH_STYLES:
            _spec_raise(
                ctx,
                f"参数 'style' 必须是 {sorted(ALLOWED_PARAGRAPH_STYLES)} 之一，"
                f"实际为 {style!r}",
            )

    if "alignment" in params:
        align = params["alignment"]
        if not isinstance(align, str) or align not in ALLOWED_ALIGNMENTS:
            _spec_raise(
                ctx,
                f"参数 'alignment' 必须是 {sorted(ALLOWED_ALIGNMENTS)} 之一，"
                f"实际为 {align!r}",
            )


def _spec_validate_add_heading(params: dict, ctx: str) -> None:
    allowed = {"text", "level"}
    unknown = set(params) - allowed
    if unknown:
        _spec_raise(
            ctx,
            f"add_heading 含未知参数 {sorted(unknown)}；允许参数: {sorted(allowed)}",
        )

    if "text" in params:
        _spec_check_type(params["text"], str, ctx, "text")

    if "level" in params:
        level = params["level"]
        if isinstance(level, bool) or not isinstance(level, int):
            _spec_raise(ctx, f"参数 'level' 必须是整数，实际为 {type(level).__name__}")
        if not (HEADING_LEVEL_MIN <= level <= HEADING_LEVEL_MAX):
            _spec_raise(
                ctx,
                f"参数 'level' 必须在 [{HEADING_LEVEL_MIN}, {HEADING_LEVEL_MAX}] 之间"
                f"（0=封面主标题 Title，1~4=一~四级章节标题），实际为 {level}",
            )


def _spec_validate_add_table(params: dict, ctx: str) -> None:
    allowed = {"data"}
    unknown = set(params) - allowed
    if unknown:
        _spec_raise(
            ctx,
            f"add_table 含未知参数 {sorted(unknown)}；仅支持参数: {sorted(allowed)}",
        )

    if "data" not in params:
        _spec_raise(ctx, "add_table 缺少必填参数 'data'")

    data = params["data"]
    if not isinstance(data, list):
        _spec_raise(ctx, f"参数 'data' 必须是二维数组，实际 {type(data).__name__}")
    if not data:
        _spec_raise(ctx, "参数 'data' 不能为空数组")
    if len(data) > MAX_TABLE_ROWS:
        _spec_raise(
            ctx,
            f"参数 'data' 行数 {len(data)} 超过上限 {MAX_TABLE_ROWS}",
        )

    for ri, row in enumerate(data):
        if not isinstance(row, list):
            _spec_raise(
                ctx,
                f"参数 'data[{ri}]' 必须是数组（一行 cells），实际 {type(row).__name__}",
            )
        if len(row) > MAX_TABLE_COLS:
            _spec_raise(
                ctx,
                f"参数 'data[{ri}]' 列数 {len(row)} 超过上限 {MAX_TABLE_COLS}",
            )
        for ci, cell in enumerate(row):
            cell_ctx = f"{ctx}.data[{ri}][{ci}]"
            if isinstance(cell, str):
                if len(cell) > MAX_CELL_TEXT_CHARS:
                    _spec_raise(
                        cell_ctx,
                        f"cell 文本长度 {len(cell)} 超过上限 "
                        f"{MAX_CELL_TEXT_CHARS} 字符",
                    )
                continue
            if isinstance(cell, dict):
                _spec_validate_run_object(
                    cell, cell_ctx, max_text_chars=MAX_CELL_TEXT_CHARS,
                )
                continue
            _spec_raise(
                cell_ctx,
                f"cell 必须是字符串或 dict（单 run 对象），实际 {type(cell).__name__}",
            )


def _spec_validate_add_page_break(params: dict, ctx: str) -> None:
    if params:
        _spec_raise(ctx, f"add_page_break 不接受任何参数，实际为 {params!r}")


# action 名称 → 校验函数；同时充当 "合法 action 集合"。
_SPEC_ACTION_VALIDATORS: dict[str, Any] = {
    "add_paragraph":  _spec_validate_add_paragraph,
    "add_heading":    _spec_validate_add_heading,
    "add_table":      _spec_validate_add_table,
    "add_page_break": _spec_validate_add_page_break,
}


def _spec_validate_command(cmd: Any, line_no: int) -> tuple[str, dict]:
    """校验单条 JSONL 命令，返回 ``(action, params)`` 便于派发器复用。"""
    ctx = f"Line {line_no}"

    if not isinstance(cmd, dict):
        _spec_raise(ctx, f"每行必须是 JSON 对象，实际 {type(cmd).__name__}")

    extra = set(cmd) - {"action", "params"}
    if extra:
        _spec_raise(ctx, f"命令仅允许 'action' / 'params' 字段，多余字段: {sorted(extra)}")

    if "action" not in cmd:
        _spec_raise(ctx, "缺少必填字段 'action'")

    action = cmd["action"]
    if not isinstance(action, str):
        _spec_raise(ctx, f"'action' 必须是字符串，实际 {type(action).__name__}")
    if action not in _SPEC_ACTION_VALIDATORS:
        _spec_raise(
            ctx,
            f"未知的 action {action!r}，允许的 action: {sorted(_SPEC_ACTION_VALIDATORS)}",
        )

    params = cmd.get("params", {})
    if not isinstance(params, dict):
        _spec_raise(ctx, f"'params' 必须是 JSON 对象，实际 {type(params).__name__}")

    _SPEC_ACTION_VALIDATORS[action](params, f"{ctx} action='{action}'")
    return action, params


# ===========================================================================
# Sandboxed local IO helpers
# ===========================================================================
#
# 所有 fs 读写都限制在
# ``WECOMAGENT_READABLE_DIRS`` / ``WECOMAGENT_WRITABLE_DIRS`` 限定
# （JSON 数组：``[{"path": "/abs/dir", "label": "..."}]``）。

ENV_READABLE = "WECOMAGENT_READABLE_DIRS"
ENV_WRITABLE = "WECOMAGENT_WRITABLE_DIRS"

# 读入 / 写出文件的大小硬上限：30 MiB。
# - 读入：避免一次性把巨型 JSONL 拉进内存撑爆进程；
# - 写出：避免生成过大的 .docx 写入磁盘（base64 后体积更大）。
MAX_FILE_SIZE_BYTES = 30 * 1024 * 1024


@functools.lru_cache(maxsize=None)
def _parse_roots(env_name: str) -> tuple[str, ...]:
    """Parse a JSON-array env var into a tuple of realpath roots (cached)."""
    raw = os.environ.get(env_name, "")
    if not raw:
        raise RuntimeError(f"环境变量 {env_name} 未设置或为空")
    parsed = json.loads(raw)
    if not isinstance(parsed, list):
        raise RuntimeError(
            f"{env_name} 必须是 JSON 数组，实际为 {type(parsed).__name__}"
        )
    roots: list[str] = []
    for it in parsed:
        if isinstance(it, str):
            it = json.loads(it)
        if not isinstance(it, dict):
            raise RuntimeError(
                f"{env_name} 元素必须是 dict 或 dict 的 JSON 字符串，"
                f"实际为 {type(it).__name__}"
            )
        p = it.get("path")
        if not isinstance(p, str) or not p.strip():
            raise RuntimeError(f"{env_name} 元素缺少有效的 path 字段: {it!r}")
        roots.append(os.path.realpath(p.strip()))
    if not roots:
        raise RuntimeError(f"环境变量 {env_name} 解析后为空")
    return tuple(roots)


def _reject_relative_segments(path: str) -> None:
    """Reject path strings that include ``.`` or ``..`` segments such as
    ``./foo``, ``../bar`` or ``/abs/path/../x``.

    Although ``os.path.realpath`` would silently normalize these away,
    accepting them would bypass the contract that callers must hand in
    explicit, fully-qualified paths inside the sandboxed roots — and
    could be abused to escape the intended directory in edge cases where
    symlinks are present.
    """
    if not path:
        return
    for seg in path.replace("\\", "/").split("/"):
        if seg in (".", ".."):
            raise ValueError(
                f"路径不允许包含 './' 或 '../' 这类相对路径片段: {path!r}"
            )


def _ensure_within(path: str, env_name: str) -> str:
    """Realpath ``path`` and ensure it lies within one of ``env_name``'s roots."""
    if not path:
        raise ValueError("path 不能为空")
    _reject_relative_segments(path)
    real = os.path.realpath(path)
    roots = _parse_roots(env_name)
    for root in roots:
        try:
            common = os.path.commonpath([real, root])
        except ValueError:
            continue
        if common == root:
            return real
    raise PermissionError(
        f"路径越权: {real} 不在 {env_name} 范围 {roots} 之内"
    )


def _read_text(path: str) -> str:
    """Read a UTF-8 text file from an allowed readable directory.

    最多读取 ``MAX_FILE_SIZE_BYTES + 1`` 字节，以便在不把超大文件
    整体载入内存的前提下判断是否超限。
    """
    real = _ensure_within(path, ENV_READABLE)
    with open(real, "rb") as f:
        data = f.read(MAX_FILE_SIZE_BYTES + 1)
    if len(data) > MAX_FILE_SIZE_BYTES:
        raise ValueError(
            f"输入文件过大：{path!r} "
            f"超过上限 {MAX_FILE_SIZE_BYTES} 字节（30 MiB）"
        )
    return data.decode("utf-8")


def _write_b64(path: str, data_b64: str, overwrite: bool = False) -> None:
    """Write a base64-encoded binary blob to an allowed writable directory.

    这里在写入前对路径做 ``os.path.islink`` 检查并显式拒绝：
      - 检查 ``path``（原始入参）：拦截 "目标位置本身就是软链" 的常见情况；
      - 检查 ``real``（realpath 结果）：作为防御纵深，覆盖悬挂软链 /
        竞态等 realpath 仍可能返回软链的边缘情况。
    """
    real = _ensure_within(path, ENV_WRITABLE)
    if os.path.islink(path) or os.path.islink(real):
        raise PermissionError(
            f"拒绝写入符号链接以避免跨目录覆盖: {path!r}"
        )
    data = base64.b64decode(data_b64, validate=True)
    if len(data) > MAX_FILE_SIZE_BYTES:
        raise ValueError(
            f"输出文件过大：解码后 {len(data)} 字节，"
            f"超过上限 {MAX_FILE_SIZE_BYTES} 字节（30 MiB）"
        )

    parent = os.path.dirname(real)
    os.makedirs(parent, exist_ok=True)

    # 创建父目录后再次解析路径，防止目录在检查和写入之间变为软链。
    real = _ensure_within(path, ENV_WRITABLE)
    if os.path.islink(path) or os.path.islink(real):
        raise PermissionError(
            f"拒绝写入符号链接以避免跨目录覆盖: {path!r}"
        )

    mode = "wb" if overwrite else "xb"
    with open(real, mode) as f:
        f.write(data)


# ===========================================================================
# Generic OOXML helpers
# ===========================================================================
#
# 注：颜色 / 数值 / 取值合法性已在上游 _spec_validate_command 阶段校验，
# 此处不再重复检查；下游 helpers 仅负责生成 OOXML 元素。


def _hex_to_rgb(color_hex: str) -> RGBColor:
    s = color_hex.lstrip("#")
    return RGBColor(int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16))


def _set_unique_child(parent, tag, new_el, insert_before=()) -> None:
    """Replace any existing ``tag`` children of ``parent`` with ``new_el``,
    inserting ahead of the first sibling listed in ``insert_before`` (the
    XSD ordering constraint). Falls back to append."""
    for existing in parent.findall(tag):
        parent.remove(existing)
    for sibling_tag in insert_before:
        sibling = parent.find(sibling_tag)
        if sibling is not None:
            sibling.addprevious(new_el)
            return
    parent.append(new_el)


def _set_east_asia_font(rPr, font_name: str) -> None:
    """Set ``w:eastAsia`` on the rFonts child of ``rPr`` (creating it if needed)."""
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def _strip_theme_color(element) -> None:
    """Remove ``themeColor``/``themeTint``/``themeShade`` from every
    ``w:color`` under ``element``.

    Built-in heading styles ship with theme-tinted colors that many
    renderers prefer over the explicit ``w:val``, leaking the accent
    color (typically blue) instead of the requested RGB.
    """
    color_tag = qn("w:color")
    color_attrs = (qn("w:themeColor"), qn("w:themeTint"), qn("w:themeShade"))
    for color_el in element.iter(color_tag):
        for key in color_attrs:
            if key in color_el.attrib:
                del color_el.attrib[key]


def _force_color_on_rpr(rPr, color_hex: str) -> None:
    """Replace any ``<w:color>`` under ``rPr`` with a plain ``w:val`` one
    (no theme attributes)."""
    for existing in rPr.findall(qn("w:color")):
        rPr.remove(existing)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color_hex.lstrip("#").upper())
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is not None:
        rFonts.addnext(color_el)
    else:
        rPr.insert(0, color_el)


def _apply_run_format(run, spec: dict) -> None:
    """Apply formatting from a run-spec dict to a python-docx Run."""
    if spec.get("bold"):
        run.bold = True
    if spec.get("italic"):
        run.italic = True
    if spec.get("underline"):
        run.underline = True
    if "color_hex" in spec:
        run.font.color.rgb = _hex_to_rgb(spec["color_hex"])
    if "size_pt" in spec:
        run.font.size = Pt(spec["size_pt"])
    if "font" in spec:
        run.font.name = spec["font"]
    if "east_asia_font" in spec:
        _set_east_asia_font(run._element.get_or_add_rPr(), spec["east_asia_font"])


def _make_borders_el(wrapper_tag: str, sides: tuple, color_hex: str, size: int):
    """Build a ``<w:tblBorders>`` / ``<w:tcBorders>`` element with all
    sides sharing the same single-line style, size and color."""
    inner = "".join(
        f'<w:{s} w:val="single" w:sz="{size}" w:color="{color_hex}"/>'
        for s in sides
    )
    return parse_xml(f'<w:{wrapper_tag} {nsdecls("w")}>{inner}</w:{wrapper_tag}>')


# ===========================================================================
# DocxBuilder — every public method (no leading underscore) is a JSONL action
# ===========================================================================


class DocxBuilder:
    """Each public method (no leading underscore) is callable as an `action`.

    所有方法不再做内部参数校验，调用方（``_dispatch``）保证传入的参数
    已通过上游 ``_spec_validate_command`` 检查。
    """

    def __init__(self) -> None:
        self.doc: Any = None  # python-docx Document

    # -- 0. Default initialization -----------------------------------------

    def _init_defaults(self) -> None:
        """Run document + page + Normal + heading defaults. Called once by
        ``run_jsonl`` before any user action; spec-level setup_* overrides win."""
        self._create_document()
        self.setup_page()
        self.setup_normal_style()
        for preset in _DEFAULT_HEADINGS:
            self.setup_heading_style(
                style_name=preset.style_name,
                size_pt=preset.size_pt,
                color_hex=preset.color_hex,
                alignment=preset.alignment,
            )

    # -- 1. Document lifecycle --------------------------------------------

    def _create_document(self) -> None:
        # Underscore-prefixed: not exposed as a JSONL action — calling it
        # twice would replace ``self.doc`` and drop everything written so far.
        self.doc = Document()
        settings = self.doc.settings.element
        zoom = settings.find(qn("w:zoom"))
        if zoom is not None and zoom.get(qn("w:percent")) is None:
            zoom.set(qn("w:percent"), "100")

    def save(self, path: str) -> None:
        """Persist the document to the local filesystem.

        ``overwrite=False`` enforces the "never clobber an existing .docx"
        guarantee that ``_pick_output_path`` makes when picking the filename.

        在编码 / 写入之前校验序列化后的 docx 体积不得超过
        ``MAX_FILE_SIZE_BYTES``；超限直接抛 ``ValueError`` 中止保存。
        """
        buf = io.BytesIO()
        self.doc.save(buf)
        size = buf.tell()
        if size > MAX_FILE_SIZE_BYTES:
            raise ValueError(
                f"输出文件过大：序列化后 {size} 字节，"
                f"超过上限 {MAX_FILE_SIZE_BYTES} 字节（30 MiB）"
            )
        data_b64 = base64.b64encode(buf.getvalue()).decode("ascii")
        _write_b64(path, data_b64, overwrite=False)

    # -- 2. Page setup ----------------------------------------------------

    def setup_page(
        self,
        width_cm: float = 21,
        height_cm: float = 29.7,
        top_cm: float = 2.4,
        bottom_cm: float = 2.4,
        left_cm: float = 2.5,
        right_cm: float = 2.5,
        header_cm: float = 1.27,
        footer_cm: float = 1.27,
    ) -> None:
        section = self.doc.sections[0]
        section.page_width = Cm(width_cm)
        section.page_height = Cm(height_cm)
        section.top_margin = Cm(top_cm)
        section.bottom_margin = Cm(bottom_cm)
        section.left_margin = Cm(left_cm)
        section.right_margin = Cm(right_cm)
        section.header_distance = Cm(header_cm)
        section.footer_distance = Cm(footer_cm)

    # -- 3. Style setup ---------------------------------------------------

    def setup_normal_style(
        self,
        font: str = "Arial",
        east_asia_font: str = "微软雅黑",
        size_pt: float = 11,
        color_hex: str = "333333",
        before_dxa: int = 60,
        after_dxa: int = 60,
        line_dxa: int = DEFAULT_LINE_DXA_NORMAL,
        widow_control: bool = False,
        snap_to_grid: bool = False,
    ) -> None:
        style = self.doc.styles["Normal"]
        style.font.name = font
        style.font.size = Pt(size_pt)
        style.font.color.rgb = _hex_to_rgb(color_hex)

        rPr = style.element.get_or_add_rPr()
        _set_east_asia_font(rPr, east_asia_font)

        style.paragraph_format.widow_control = widow_control

        pPr = style.element.get_or_add_pPr()

        snap = OxmlElement("w:snapToGrid")
        snap.set(qn("w:val"), "1" if snap_to_grid else "0")
        _set_unique_child(pPr, qn("w:snapToGrid"), snap, _PPR_SPACING_ANCHORS)

        sp = self._make_spacing_el(before_dxa, after_dxa, line_dxa)
        _set_unique_child(pPr, qn("w:spacing"), sp, _PPR_SPACING_ANCHORS)

    def setup_heading_style(
        self,
        style_name: str,
        size_pt: float,
        color_hex: str = "1A1A1A",
        bold: bool = True,
        font: str = "Arial",
        line_dxa: int = DEFAULT_LINE_DXA_HEADING,
        before_dxa: int = 0,
        after_dxa: int = 0,
        alignment: str | None = None,
        keep_with_next: bool = True,
        keep_together: bool = True,
    ) -> None:
        """Configure Title / Subtitle / Heading 1..N styles.

        仅由 ``_init_defaults`` 内部调用，传入的参数来自 ``_DEFAULT_HEADINGS``
        预置常量，已知合法；不再做单独校验。
        """
        style = self.doc.styles[style_name]
        style.font.name = font
        style.font.size = Pt(size_pt)
        style.font.bold = bold
        style.font.color.rgb = _hex_to_rgb(color_hex)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.keep_with_next = keep_with_next
        style.paragraph_format.keep_together = keep_together
        if alignment is not None:
            style.paragraph_format.alignment = _PARAGRAPH_ALIGN[alignment]

        pPr = style.element.get_or_add_pPr()
        sp = self._make_spacing_el(before_dxa, after_dxa, line_dxa)
        _set_unique_child(pPr, qn("w:spacing"), sp, _PPR_SPACING_ANCHORS)

        # Drop the decorative theme-accent bottom border that built-in
        # Title (and a few headings) ship with.
        for pBdr in pPr.findall(qn("w:pBdr")):
            pPr.remove(pBdr)

        # Strip themeColor/Tint/Shade everywhere in the style element so
        # the requested RGB actually wins on renderers that prefer theme.
        _strip_theme_color(style.element)

        # Mirror the color onto the paired *character* style (Title↔TitleChar,
        # Heading 1↔Heading1Char, …): some renderers (WPS, Word for Mac /
        # Online, the WeCom doc preview) apply the linked char style's rPr
        # to runs in preference to the paragraph style's rPr.
        self._sync_linked_char_style_color(style, color_hex)

        # Strip leaked python-docx template defaults and add bCs/szCs that
        # the high-level setters skip; do the same on the linked char style.
        self._normalize_heading_style_element(style.element, bold)
        linked_char = self._resolve_linked_char_style_element(style)
        if linked_char is not None:
            self._normalize_heading_style_element(linked_char, bold)

    @staticmethod
    def _make_spacing_el(before_dxa: int, after_dxa: int, line_dxa: int):
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:before"), str(before_dxa))
        sp.set(qn("w:after"), str(after_dxa))
        sp.set(qn("w:line"), str(line_dxa))
        sp.set(qn("w:lineRule"), "auto")
        return sp

    def _sync_linked_char_style_color(self, paragraph_style, color_hex: str) -> None:
        """Mirror ``color_hex`` onto the paragraph style's linked char style
        (silent no-op if there is no link or it cannot be resolved)."""
        target = self._resolve_linked_char_style_element(paragraph_style)
        if target is None:
            return
        rPr = target.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            target.append(rPr)
        _force_color_on_rpr(rPr, color_hex)

    def _resolve_linked_char_style_element(self, paragraph_style):
        """Return the ``<w:style>`` element of the linked char style, or None."""
        link_el = paragraph_style.element.find(qn("w:link"))
        if link_el is None:
            return None
        char_style_id = link_el.get(qn("w:val"))
        if not char_style_id:
            return None
        for s in self.doc.styles.element.findall(qn("w:style")):
            if s.get(qn("w:styleId")) == char_style_id:
                return s
        return None

    @staticmethod
    def _normalize_heading_style_element(style_element, bold: bool) -> None:
        """Strip python-docx template defaults on built-in heading styles
        and add the bCs / szCs siblings that high-level setters skip.

        Removed: pPr/contextualSpacing (Title), pPr/numPr (Subtitle),
        rPr/i, rPr/iCs, rPr/spacing, rPr/kern, and theme-bound rFonts
        attributes (asciiTheme/hAnsiTheme/eastAsiaTheme/cstheme).
        Added: rPr/bCs (paired with rPr/b) for CJK / complex-script bold;
        rPr/szCs synced to rPr/sz so font.size actually applies.
        """
        pPr = style_element.find(qn("w:pPr"))
        if pPr is not None:
            for tag in ("w:contextualSpacing", "w:numPr"):
                for child in pPr.findall(qn(tag)):
                    pPr.remove(child)

        rPr = style_element.find(qn("w:rPr"))
        if rPr is None:
            return

        for tag in ("w:i", "w:iCs", "w:spacing", "w:kern"):
            for child in rPr.findall(qn(tag)):
                rPr.remove(child)

        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is not None:
            for attr in (
                "w:asciiTheme",
                "w:hAnsiTheme",
                "w:eastAsiaTheme",
                "w:cstheme",
            ):
                key = qn(attr)
                if key in rFonts.attrib:
                    del rFonts.attrib[key]

        if bold:
            b = rPr.find(qn("w:b"))
            if b is not None and rPr.find(qn("w:bCs")) is None:
                bCs = OxmlElement("w:bCs")
                b.addnext(bCs)

        sz = rPr.find(qn("w:sz"))
        if sz is not None:
            sz_val = sz.get(qn("w:val"))
            if sz_val:
                szCs = rPr.find(qn("w:szCs"))
                if szCs is None:
                    szCs = OxmlElement("w:szCs")
                    sz.addnext(szCs)
                szCs.set(qn("w:val"), sz_val)

    # -- 4. Content blocks ------------------------------------------------

    def add_heading(self, text: str = "", level: int = 1) -> None:
        self.doc.add_heading(text, level=level)

    def add_paragraph(
        self,
        text: str | None = None,
        runs: list[dict] | None = None,
        style: str | None = None,
        alignment: str | None = None,
    ) -> None:
        """Add a paragraph.

        - ``text``  : single-run plain text.
        - ``runs``  : list of run-specs (bold/italic/color/size...). If both
                      ``text`` and ``runs`` are given, ``runs`` wins.
        - ``style`` : built-in style name, e.g. 'List Bullet', 'List Number',
                      'Subtitle'.
        """
        p = self.doc.add_paragraph(style=style) if style else self.doc.add_paragraph()
        if alignment:
            p.alignment = _PARAGRAPH_ALIGN[alignment]
        if runs:
            for r in runs:
                run = p.add_run(r.get("text", ""))
                _apply_run_format(run, r)
        elif text is not None:
            p.add_run(text)

    def add_page_break(self) -> None:
        self.doc.add_page_break()

    # -- 5. Table ---------------------------------------------------------

    def add_table(
        self,
        data: list,
        col_widths_dxa: list[int] | None = None,
        total_width_dxa: int = DEFAULT_TABLE_TOTAL_DXA,
        border_color_hex: str = DEFAULT_CELL_BORDER_COLOR_HEX,
        border_size: int = DEFAULT_CELL_BORDER_SIZE,
        cell_margin_dxa: tuple | list = (0, 108, 0, 108),  # top, left, bottom, right
        header_shading_hex: str | None = None,
        alignment: str = "center",
        cell_v_align: str = "center",
    ) -> None:
        """Add a fixed-layout table.

        ``data`` is a list of rows. Each row is a list of cells. A cell can
        be a plain string OR a dict like
        ``{"text": "...", "bold": true, "color_hex": "FF0000"}``.

        The first row is auto-tagged with ``<w:tblHeader/>`` so it repeats
        on page breaks. Header shading is OFF by default — pass
        ``header_shading_hex="2972F4"`` to opt in.
        """
        if not data:
            return

        rows = len(data)
        cols = max(len(r) for r in data)
        col_widths = self._resolve_col_widths(cols, col_widths_dxa, total_width_dxa)

        table = self._create_blank_table(rows, cols, alignment)
        self._apply_table_width(table)
        self._apply_table_borders(
            table,
            DEFAULT_TABLE_BORDER_COLOR_HEX,
            DEFAULT_TABLE_BORDER_SIZE,
        )
        self._apply_table_fixed_layout(table)
        self._strip_table_look(table)
        self._apply_table_grid(table, col_widths)
        self._mark_header_row(table)
        self._fill_table_cells(
            table, data, col_widths,
            cell_margin_dxa, header_shading_hex, cell_v_align,
            border_color_hex, border_size,
        )

    # -- 5.1 Table internals ---------------------------------------------

    @staticmethod
    def _resolve_col_widths(
        cols: int,
        col_widths_dxa: list[int] | None,
        total_width_dxa: int,
    ) -> list[int]:
        if col_widths_dxa:
            return list(col_widths_dxa)
        base = total_width_dxa // cols
        widths = [base] * cols
        widths[-1] += total_width_dxa - base * cols  # absorb rounding
        return widths

    def _create_blank_table(self, rows: int, cols: int, alignment: str):
        # Intentionally no ``table.style = "Table Grid"`` — we provide all
        # visual properties explicitly, and a built-in style would leak its
        # own border / shading defaults.
        table = self.doc.add_table(rows=rows, cols=cols)
        table.alignment = _TABLE_ALIGN.get(alignment, WD_TABLE_ALIGNMENT.CENTER)
        return table

    @staticmethod
    def _apply_table_width(table) -> None:
        # ``tblW`` declares ``auto``; the actual width is dictated by
        # ``tblLayout=fixed`` plus the explicit ``<w:gridCol>`` widths.
        tblPr = table._tbl.tblPr
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), "0")
        tblW.set(qn("w:type"), "auto")
        _set_unique_child(tblPr, qn("w:tblW"), tblW, _TBL_PR_TBLW_ANCHORS)

    @staticmethod
    def _apply_table_borders(table, color_hex: str, size: int) -> None:
        # Table-level frame only; per-cell borders are added separately so
        # the grid stays visible on renderers that ignore <w:tblBorders>.
        tblPr = table._tbl.tblPr
        el = _make_borders_el(
            "tblBorders",
            ("top", "left", "bottom", "right", "insideH", "insideV"),
            color_hex,
            size,
        )
        _set_unique_child(tblPr, qn("w:tblBorders"), el, _TBL_PR_BORDERS_ANCHORS)

    @staticmethod
    def _apply_table_fixed_layout(table) -> None:
        tblPr = table._tbl.tblPr
        el = OxmlElement("w:tblLayout")
        el.set(qn("w:type"), "fixed")
        _set_unique_child(tblPr, qn("w:tblLayout"), el, _TBL_PR_LAYOUT_ANCHORS)

    @staticmethod
    def _strip_table_look(table) -> None:
        # We don't attach a table style, so <w:tblLook> (firstRow / banding /
        # ... toggles) is inert noise relative to the target docx.
        tblPr = table._tbl.tblPr
        for el in tblPr.findall(qn("w:tblLook")):
            tblPr.remove(el)

    @staticmethod
    def _apply_table_grid(table, col_widths_dxa: list[int]) -> None:
        grid = table._tbl.find(qn("w:tblGrid"))
        if grid is None:
            return
        for col in list(grid.findall(qn("w:gridCol"))):
            grid.remove(col)
        for w in col_widths_dxa:
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(w))
            grid.append(gc)

    @staticmethod
    def _mark_header_row(table) -> None:
        # Tag the first row with <w:tblHeader/> so it repeats on page breaks.
        if not table.rows:
            return
        tr = table.rows[0]._tr
        trPr = tr.find(qn("w:trPr"))
        if trPr is None:
            trPr = OxmlElement("w:trPr")
            tr.insert(0, trPr)  # trPr precedes <w:tc> per the OOXML schema
        if trPr.find(qn("w:tblHeader")) is None:
            trPr.append(OxmlElement("w:tblHeader"))

    def _fill_table_cells(
        self,
        table,
        data: list,
        col_widths_dxa: list[int],
        cell_margin_dxa: tuple | list,
        header_shading_hex: str | None,
        cell_v_align: str,
        cell_border_color_hex: str,
        cell_border_size: int,
    ) -> None:
        cols = len(col_widths_dxa)
        for ri, row_data in enumerate(data):
            for ci in range(cols):
                cell = table.rows[ri].cells[ci]
                value = row_data[ci] if ci < len(row_data) else None

                self._set_cell_width(cell, col_widths_dxa[ci])
                self._write_cell_content(cell, value)

                # Apply tcBorders → shd → tcMar → vAlign in this order so
                # the anchor lookups in ``_set_unique_child`` resolve.
                self._apply_cell_borders(
                    cell, cell_border_color_hex, cell_border_size,
                )
                if ri == 0 and header_shading_hex:
                    self._apply_cell_shading(cell, header_shading_hex)
                self._apply_cell_margin(cell, cell_margin_dxa)
                self._apply_cell_v_align(cell, cell_v_align)

    @staticmethod
    def _set_cell_width(cell, width_dxa: int) -> None:
        # cell.width takes EMU-typed Length; convert dxa → EMU.
        cell.width = Emu(width_dxa * EMU_PER_DXA)

    @staticmethod
    def _write_cell_content(cell, value) -> None:
        # ``cell.text = ""`` leaves an empty <w:r/> placeholder that renders
        # as a stray empty run; clear runs on the first paragraph instead.
        p = cell.paragraphs[0]
        for run in list(p.runs):
            run._element.getparent().remove(run._element)
        if value is None:
            return
        if isinstance(value, dict):
            run = p.add_run(value.get("text", ""))
            _apply_run_format(run, value)
        else:
            p.add_run(str(value))

    @staticmethod
    def _apply_cell_borders(cell, color_hex: str, size: int) -> None:
        # Per-cell <w:tcBorders> in addition to the table-level frame so
        # the grid stays intact on renderers that disagree on which level
        # of border is authoritative.
        tcPr = cell._tc.get_or_add_tcPr()
        el = _make_borders_el(
            "tcBorders",
            ("top", "left", "bottom", "right"),
            color_hex,
            size,
        )
        _set_unique_child(tcPr, qn("w:tcBorders"), el, _TC_PR_BORDERS_ANCHORS)

    @staticmethod
    def _apply_cell_shading(cell, fill_hex: str) -> None:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(
            f'<w:shd {nsdecls("w")} w:val="clear" '
            f'w:color="auto" w:fill="{fill_hex}"/>'
        )
        _set_unique_child(tcPr, qn("w:shd"), shd, _TC_PR_SHD_ANCHORS)

    @staticmethod
    def _apply_cell_margin(cell, margin_dxa: tuple | list) -> None:
        # margin_dxa = (top, left, bottom, right).
        top, left, bottom, right = margin_dxa
        tcPr = cell._tc.get_or_add_tcPr()
        el = parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'  <w:top w:w="{top}" w:type="dxa"/>'
            f'  <w:left w:w="{left}" w:type="dxa"/>'
            f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
            f'  <w:right w:w="{right}" w:type="dxa"/>'
            f'</w:tcMar>'
        )
        _set_unique_child(tcPr, qn("w:tcMar"), el, _TC_PR_MAR_ANCHORS)

    @staticmethod
    def _apply_cell_v_align(cell, val: str) -> None:
        tcPr = cell._tc.get_or_add_tcPr()
        el = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{val}"/>')
        _set_unique_child(tcPr, qn("w:vAlign"), el, _TC_PR_VALIGN_ANCHORS)


# ===========================================================================
# JSONL spec parsing
# ===========================================================================


def _iter_spec_commands(spec_path: str) -> Iterator[tuple[int, Any]]:
    """Yield ``(line_no, cmd)`` from a JSONL spec file.

    Tolerates blank lines, leading UTF-8 BOM, ``//`` / ``#`` comment lines,
    and JSON objects pretty-printed across multiple lines (uses
    ``raw_decode`` to consume one object at a time). ``line_no`` is the
    1-based line where each object *starts*.
    """
    text = _read_text(spec_path)

    if text.startswith("\ufeff"):
        text = text[1:]

    decoder = json.JSONDecoder()
    idx = 0
    n = len(text)

    # Incremental newline counter: scans only the disjoint segment
    # text[line_cursor:idx] each iteration → O(N) total.
    line_cursor = 0
    line_no = 1

    while idx < n:
        ch = text[idx]

        if ch.isspace():
            idx += 1
            continue

        if ch == "#" or text.startswith("//", idx):
            nl = text.find("\n", idx)
            if nl == -1:
                break
            idx = nl + 1
            continue

        line_no += text.count("\n", line_cursor, idx)
        line_cursor = idx
        start_line = line_no

        try:
            cmd, end = decoder.raw_decode(text, idx)
        except json.JSONDecodeError as e:
            raise SpecTypeError(f"Line {start_line}: invalid JSON: {e}") from e

        yield start_line, cmd
        idx = end


# ===========================================================================
# Dispatcher & runner
# ===========================================================================


def _dispatch(builder: DocxBuilder, action: str, params: dict) -> None:
    """执行单条已通过校验的命令。

    上游 ``_spec_validate_command`` 已确保 ``action`` 合法且 ``params``
    形态正确，这里直接派发到对应方法即可。
    """
    method = getattr(builder, action)
    method(**params)


def run_jsonl(spec_path: str, output: str) -> str:
    if not output:
        raise SpecTypeError("An output path must be provided.")

    # 1) 一次性把整份 JSONL 读出来，全部走完上游校验，再开始写文档。
    #    任何参数问题都会以 SpecTypeError("类型错误，无法执行") 抛出。
    #    命令总数受 MAX_COMMANDS 限制，超限立即中止以避免下游构建阶段
    #    因海量命令而 OOM / 卡死。
    commands: list[tuple[str, dict]] = []
    for line_no, cmd in _iter_spec_commands(spec_path):
        if len(commands) >= MAX_COMMANDS:
            raise SpecTypeError(
                f"Line {line_no}: 命令总数超过上限 {MAX_COMMANDS}"
            )
        action, params = _spec_validate_command(cmd, line_no)
        commands.append((action, params))

    # 2) 校验通过 → 实际生成文档并保存。
    builder = DocxBuilder()
    builder._init_defaults()
    for action, params in commands:
        _dispatch(builder, action, params)
    builder.save(output)

    return output


# ===========================================================================
# CLI
# ===========================================================================


def _pick_output_path(spec_path: str) -> str:
    """Pick a non-conflicting ``.docx`` path under the writable root.

    Filename derives from the spec's stem (``report.jsonl`` → ``report.docx``);
    on conflict a timestamp suffix is appended to avoid overwriting.
    """
    target_dir = os.path.join(_parse_roots(ENV_WRITABLE)[0], "docx")
    target_dir = _ensure_within(target_dir, ENV_WRITABLE)
    stem = Path(spec_path).stem or "document"
    if not re.fullmatch(r"[A-Za-z0-9_.\-]{1,128}", stem):
        stem = "document"

    candidate = os.path.join(target_dir, f"{stem}.docx")
    if not os.path.lexists(candidate):
        return candidate
    # Multi-user host: combine millisecond timestamp with PID to avoid
    # collisions between concurrent processes within the same millisecond.
    suffix = f"{int(time.time_ns() // 1_000_000)}_{os.getpid()}"
    candidate = os.path.join(target_dir, f"{stem}_{suffix}.docx")
    if not os.path.lexists(candidate):
        return candidate
    # 时间戳 + PID 仍然冲突属于极端异常情况，直接报错而非覆盖既有文件。
    raise FileExistsError(f"无法生成唯一的输出路径：{candidate} 已存在")


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Build a .docx file from a JSONL spec."
    )
    parser.add_argument("spec", help="Path to the JSONL spec file")
    args = parser.parse_args()

    try:
        output = _pick_output_path(args.spec)
    except Exception as e:
        print("Error: failed to pick output path")
        sys.exit(2)

    try:
        saved = run_jsonl(args.spec, output=output)
    except TypeError as e:
        # 上游 JSONL 校验抛出的 SpecTypeError(继承 TypeError)，统一
        # 转译成 "类型错误，无法执行" 提示。
        print(f"Error: 类型错误，无法执行: {e}", file=sys.stderr)
        sys.exit(2)
    except PermissionError:
        print("Error: 路径不在允许范围内", file=sys.stderr)
        sys.exit(2)
    except Exception:
        print("Error: 执行失败，请检查输入文件格式或稍后重试", file=sys.stderr)
        sys.exit(2)

    print(f"Successfully built {saved}")


if __name__ == "__main__":
    main()
